"""
word_export_railway.py - 軌二部 Word 報告輸出模組

將檢核結果輸出為 Word 表格報告，格式參照業務部標準格式：
- 5欄橫式表格：項次、檢查項目、條款、條款摘要、備註
- 標題列：藍色背景
- 大項目：綠色背景
- 子項目：白色背景
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_ALIGN_VERTICAL
from datetime import datetime
import json


class RailwayWordExporter:
    """軌二部 Word 報告生成器"""

    # 顏色定義（參照業務部格式）
    COLOR_HEADER = 'A5C9EB'      # 標題列：淺藍色
    COLOR_MAIN_ITEM = 'C1F0C7'   # 大項目：淺綠色
    COLOR_SUB_ITEM = 'FFFFFF'    # 子項目：白色

    def __init__(self):
        self.doc = Document()
        self._setup_document_styles()
        self.table = None

    def _setup_document_styles(self):
        """設定文件樣式"""
        # 設定預設字體為標楷體
        self.doc.styles['Normal'].font.name = '標楷體'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        self.doc.styles['Normal'].font.size = Pt(12)

    def _set_cell_color(self, cell, color_hex):
        """設定儲存格背景色

        Args:
            cell: 儲存格物件
            color_hex: 顏色16進位碼（如 'A5C9EB'）
        """
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color_hex)
        cell._element.get_or_add_tcPr().append(shading_elm)

    def _set_cell_border(self, cell):
        """設定儲存格邊框"""
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()

        # 定義邊框樣式
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # 邊框寬度
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcPr.append(border)

    def add_title(self, project_name="契約審查報告", department="軌二部"):
        """添加報告標題"""
        title = self.doc.add_heading(project_name, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加資訊段落
        info = self.doc.add_paragraph()
        info.add_run(f"部門：{department}\n")
        info.add_run(f"日期：{datetime.now().strftime('%Y年%m月%d日')}\n")
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph()  # 空行

    def add_cover_page(self, review_comments, project_number="", project_name=""):
        """添加封面頁（參照業務部格式）

        Args:
            review_comments: 審查意見列表（違反備註條件的提醒）
            project_number: 業務編號（選填）
            project_name: 業務簡稱（選填）
        """
        # 創建封面表格（4列4欄）
        cover_table = self.doc.add_table(rows=4, cols=4)
        cover_table.style = 'Table Grid'

        # === 第1列：業務資訊 ===
        row0 = cover_table.rows[0]
        row0.cells[0].text = '業務編號'
        row0.cells[1].text = project_number
        row0.cells[2].text = '業務簡稱'
        row0.cells[3].text = project_name

        # 設定樣式
        for cell in row0.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self._set_cell_border(cell)

        # === 第2列：審查資訊 ===
        row1 = cover_table.rows[1]
        row1.cells[0].text = '審查日期'
        row1.cells[1].text = datetime.now().strftime('%Y年%m月%d日')
        row1.cells[2].text = '審查結論'
        row1.cells[3].text = ''  # 可由用戶自行填寫

        for cell in row1.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self._set_cell_border(cell)

        # === 第3列：審查意見（業務部） ===
        row2 = cover_table.rows[2]

        # 合併第2-4欄
        merged_cell = row2.cells[1].merge(row2.cells[2]).merge(row2.cells[3])

        row2.cells[0].text = '業務部'
        row2.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row2.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        self._set_cell_border(row2.cells[0])

        # 填入審查意見
        if review_comments:
            merged_cell.text = '\n'.join([f"{i+1}. {comment}" for i, comment in enumerate(review_comments)])
        else:
            merged_cell.text = '（無重點提醒事項）'

        merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        self._set_cell_border(merged_cell)

        # === 第4列：法律部 ===
        row3 = cover_table.rows[3]
        row3.cells[0].text = '法律部'

        # 合併第2-4欄
        merged_cell_law = row3.cells[1].merge(row3.cells[2]).merge(row3.cells[3])
        merged_cell_law.text = '審查人員 簽名：'

        row3.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row3.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        self._set_cell_border(row3.cells[0])

        merged_cell_law.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        merged_cell_law.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        self._set_cell_border(merged_cell_law)

        # 添加分頁符
        self.doc.add_page_break()

    def create_main_table(self):
        """創建主表格（參照業務部格式）"""
        # 創建表格（初始1列，後續動態添加）
        self.table = self.doc.add_table(rows=1, cols=5)
        self.table.style = 'Table Grid'

        # 設定欄寬
        self.table.columns[0].width = Cm(1.5)   # 項次
        self.table.columns[1].width = Cm(4.0)   # 檢查項目
        self.table.columns[2].width = Cm(2.5)   # 條款
        self.table.columns[3].width = Cm(5.0)   # 條款摘要
        self.table.columns[4].width = Cm(3.0)   # 備註

        # 設定標題列
        header_row = self.table.rows[0]
        headers = ['項次', '檢查項目', '條款', '條款摘要', '備註']

        for i, header_text in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header_text

            # 設定標題列樣式
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # 設定背景色和邊框
            self._set_cell_color(cell, self.COLOR_HEADER)
            self._set_cell_border(cell)

    def add_main_item_row(self, item_number, item_description):
        """添加大項目列（綠色背景）

        Args:
            item_number: 主項次（如 "1."）
            item_description: 主項說明
        """
        row = self.table.add_row()

        # 填寫項次和檢查項目
        row.cells[0].text = item_number
        row.cells[1].text = item_description

        # 設定樣式
        for i, cell in enumerate(row.cells):
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if i > 0 else WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # 設定綠色背景
            self._set_cell_color(cell, self.COLOR_MAIN_ITEM)
            self._set_cell_border(cell)

            # 項次加粗
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True

    def add_sub_item_row(self, item_data):
        """添加子項目列（白色背景）

        Args:
            item_data (dict): 項目資料 {item_number, check_item, clause, clause_summary, remark}
        """
        row = self.table.add_row()

        # 填寫各欄位
        row.cells[0].text = item_data.get('item_number', '')
        row.cells[1].text = item_data.get('check_item', '')
        row.cells[2].text = item_data.get('clause', '')
        row.cells[3].text = item_data.get('clause_summary', '')
        row.cells[4].text = item_data.get('remark', '')

        # 設定樣式
        for i, cell in enumerate(row.cells):
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if i > 0 else WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            # 設定白色背景和邊框
            self._set_cell_color(cell, self.COLOR_SUB_ITEM)
            self._set_cell_border(cell)

    def export_results(self, results, json_path, output_path, project_name="契約審查報告",
                      project_number="", project_brief="", include_cover=True):
        """匯出檢核結果為 Word 報告

        Args:
            results (list or dict): 檢核結果
            json_path (str): JSON 檔案路徑（用於讀取完整結構）
            output_path (str): 輸出檔案路徑
            project_name (str): 專案名稱
            project_number (str): 業務編號（選填）
            project_brief (str): 業務簡稱（選填）
            include_cover (bool): 是否包含封面頁（業務部=True，軌二部=False）
        """
        # 處理結果為字典（項次 -> 結果）
        results_dict = {}
        if isinstance(results, dict):
            results = [results]

        # 收集所有審查意見（違反備註條件的提醒）
        review_comments = []
        for result in results:
            item_num = result.get('item_number', '')
            results_dict[item_num] = result

            # 提取審查意見
            review_comment = result.get('review_comment', '')
            if review_comment:
                check_item = result.get('check_item', '')
                review_comments.append(f"【{item_num} {check_item}】{review_comment}")

        # 根據部門決定是否添加封面頁
        if include_cover:
            # 業務部：添加封面頁（含審查意見）
            self.add_cover_page(review_comments, project_number, project_brief)
        else:
            # 軌二部：直接添加標題，不需要封面頁
            self.add_title(project_name=project_name, department="軌二部")

        # 創建主表格
        self.create_main_table()

        # 讀取完整 JSON 結構
        with open(json_path, 'r', encoding='utf-8') as f:
            checklist_data = json.load(f)

        # 遍歷 JSON 結構，按順序添加
        for main_item in checklist_data:
            main_num = main_item.get('主項次', '')
            main_desc = main_item.get('主項說明', '')

            # 添加大項目列
            self.add_main_item_row(f"{main_num}.", main_desc)

            # 添加子項目
            for sub_item in main_item.get('子項目', []):
                sub_num = sub_item.get('項次', '')

                # 檢查是否有更深層的子項目
                has_sub_sub = '子項目' in sub_item and sub_item['子項目']

                if has_sub_sub:
                    # 有子子項目，遍歷子子項目
                    for sub_sub_item in sub_item['子項目']:
                        sub_sub_num = sub_sub_item.get('項次', '')

                        # 從結果中獲取分析資料
                        if sub_sub_num in results_dict:
                            item_data = self._parse_result(results_dict[sub_sub_num])
                        else:
                            # 沒有結果，使用空白
                            item_data = {
                                'item_number': sub_sub_num,
                                'check_item': sub_sub_item.get('檢查項目', ''),
                                'clause': sub_sub_item.get('條款', ''),
                                'clause_summary': sub_sub_item.get('條款摘要', ''),
                                'remark': sub_sub_item.get('備註', '')
                            }

                        self.add_sub_item_row(item_data)
                else:
                    if sub_num in results_dict:
                        item_data = self._parse_result(results_dict[sub_num])
                    else:
                        item_data = {
                            'item_number': sub_num,
                            'check_item': sub_item.get('檢查項目', ''),
                            'clause': sub_item.get('條款', ''),
                            'clause_summary': sub_item.get('條款摘要', ''),
                            'remark': sub_item.get('備註', '')
                        }

                    self.add_sub_item_row(item_data)

        # 儲存文件
        self.doc.save(output_path)
        print(f"✅ Word 報告已輸出至: {output_path}")

    def _parse_result(self, result):
        """解析檢核結果，提取表格所需資料

        Args:
            result (dict): 檢核結果，包含 analysis 欄位

        Returns:
            dict: 表格資料 {item_number, check_item, clause, clause_summary, remark}
        """
        item_number = result.get('item_number', '')
        check_item = result.get('check_item', '')

        # 解析 LLM 分析結果
        analysis = result.get('analysis', '')

        clause = ''
        clause_summary = ''
        remark = ''

        # 從分析結果中提取各欄位
        lines = analysis.split('\n')
        current_section = None

        for line in lines:
            line_stripped = line.strip()
            if line_stripped.startswith('條款：') or line_stripped.startswith('條款:'):
                clause = line_stripped.replace('條款：', '').replace('條款:', '').strip()
                current_section = 'clause'
            elif line_stripped.startswith('條款摘要：') or line_stripped.startswith('條款摘要:'):
                clause_summary = line_stripped.replace('條款摘要：', '').replace('條款摘要:', '').strip()
                current_section = 'summary'
            elif line_stripped.startswith('備註：') or line_stripped.startswith('備註:'):
                remark = line_stripped.replace('備註：', '').replace('備註:', '').strip()
                current_section = 'remark'
            elif line_stripped and current_section:
                # 續行內容
                if current_section == 'summary':
                    clause_summary += '\n' + line_stripped
                elif current_section == 'remark':
                    remark += '\n' + line_stripped

        return {
            'item_number': item_number,
            'check_item': check_item,
            'clause': clause,
            'clause_summary': clause_summary,
            'remark': remark
        }


def export_to_word(results, json_path, output_path, project_name="契約審查報告",
                  project_number="", project_brief="", include_cover=True):
    """快速匯出函數

    Args:
        results: 檢核結果（單一 dict 或 list of dicts）
        json_path: JSON 檔案路徑
        output_path: 輸出檔案路徑
        project_name: 專案名稱
        project_number: 業務編號（選填）
        project_brief: 業務簡稱（選填）
        include_cover: 是否包含封面頁（業務部=True，軌二部=False）
    """
    exporter = RailwayWordExporter()
    exporter.export_results(results, json_path, output_path, project_name,
                          project_number, project_brief, include_cover)


if __name__ == "__main__":
    # 測試範例
    test_results = [
        {
            'item_number': '1.1',
            'check_item': '給付條件',
            'analysis': '''條款：契約文件第5條
        條款摘要：■無□有
        備註：依據契約文件第5條規定，本契約採總包價法，未保留後續擴充權利，因此勾選「■無」。'''
        }
    ]

    test_json_path = r"C:\Users\user\Desktop\中興ISO案\contract_checker_app\output\test.json"
    export_to_word(test_results, test_json_path, "test_output.docx", "測試專案契約審查報告")
